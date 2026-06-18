"""
Generates the Customer Segmentation project planning documents as Word .docx files.
Each document captures one topic from our planning conversation.
Run:  python3 _generate_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # dark blue
GREY = RGBColor(0x55, 0x55, 0x55)


def new_doc(title, subtitle):
    doc = Document()
    # base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_heading(title, level=0)
    for run in t.runs:
        run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    r = sub.add_run(subtitle)
    r.italic = True
    r.font.color.rgb = GREY
    r.font.size = Pt(11)

    meta = doc.add_paragraph()
    m = meta.add_run("Customer Segmentation Project  |  Planning Notes  |  17 June 2026")
    m.font.size = Pt(9)
    m.font.color.rgb = GREY
    doc.add_paragraph("")
    return doc


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def para(doc, text):
    return doc.add_paragraph(text)


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph("")
    return t


def save(doc, name):
    path = os.path.join(OUT_DIR, name)
    doc.save(path)
    print("wrote", name)


# ---------------------------------------------------------------------------
# 00 - PROJECT OVERVIEW (index)
# ---------------------------------------------------------------------------
d = new_doc("Project Overview", "What we are building and why")
h1(d, "One-line summary")
para(d, "A rigorous, explainable customer-segmentation project built on real e-commerce "
        "transaction data (UCI 'Online Retail II'). It segments customers into a small number "
        "of statistically validated groups, predicts each customer's future value with a "
        "probabilistic model, and translates the result into concrete marketing actions.")

h1(d, "The interview-ready story")
para(d, "\"I took 1M+ real e-commerce transactions, cleaned them into customer-level records, "
        "engineered RFM features, and segmented customers using K-Means - but instead of trusting "
        "the elbow method blindly, I validated the number of clusters with silhouette and gap "
        "statistics, tested cluster stability with bootstrapping, and benchmarked K-Means against "
        "Gaussian Mixture and hierarchical clustering. I then replaced the naive lifetime-value "
        "formula with a probabilistic BG/NBD + Gamma-Gamma model to predict each customer's future "
        "value, and translated the segments into concrete marketing actions.\"")

h1(d, "Project decisions (locked in)")
table(d, ["Decision", "Choice"], [
    ["Goal", "Portfolio / CV piece"],
    ["Dataset", "Real - UCI 'Online Retail II' (downloaded & verified)"],
    ["Statistical rigor", "Full - validation, method comparison, probabilistic CLV"],
    ["Format of work", "GitHub repo, numbered notebooks, reusable src/ modules"],
    ["Working style", "Phase by phase; everything must be explainable"],
])

h1(d, "Document index")
para(d, "This folder contains the planning conversation, split by topic:")
table(d, ["Document", "Contents"], [
    ["00_Project_Overview", "This file - summary, decisions, index"],
    ["01_Source_Notebook_Analysis", "Analysis of the original synthetic-data notebook and its weaknesses"],
    ["02_Project_Value", "Business value + career value; why statistical rigor matters"],
    ["03_Dataset", "Online Retail II - verified facts, columns, data-quality quirks"],
    ["04_Column_Importance_Map", "Role and importance of every column; the feature-engineering principle"],
    ["05_Problem_Formulation", "Formal problem statement, objectives, hypotheses, expected results"],
    ["06_Repo_Structure_and_Roadmap", "Planned repo layout, notebooks, methods map, data flow"],
])
save(d, "00_Project_Overview.docx")


# ---------------------------------------------------------------------------
# 01 - SOURCE NOTEBOOK ANALYSIS
# ---------------------------------------------------------------------------
d = new_doc("Source Notebook Analysis", "Understanding the original synthetic-data notebook")
h1(d, "What the original notebook was")
para(d, "The project started from an existing 58-page notebook titled "
        "'Marketing_Analytics_&_Customer_Segmentation'. It ran an end-to-end pipeline in Python "
        "(pandas, numpy, matplotlib, seaborn, scikit-learn) on a SYNTHETIC dataset of 125,000 "
        "customers x 14 columns (marketing_customer.csv).")

h1(d, "The pipeline it followed")
numbered(d, "Cleaning: impute missing categoricals with 'unknown', numerics with median; fix dtypes; "
            "IQR outlier capping (winsorization); drop 5,000 duplicate rows; drop a junk text column.")
numbered(d, "EDA: histograms, boxplots, categorical counts, scatterplots, revenue by channel/location, "
            "correlation heatmap (all correlations near zero - by design).")
numbered(d, "Feature engineering: RFM (Recency, Frequency, Monetary), 1-5 RFM scores, rule-based "
            "segments (Champions/Loyal/Potential/At-Risk/Lost), and a deterministic CLV formula.")
numbered(d, "Segmentation: K-Means (elbow -> K=4, scaled features, PCA viz) and DBSCAN (collapsed to "
            "one blob - correctly judged unsuitable).")
numbered(d, "Insights & recommendations: marketing actions per segment.")

h1(d, "Key weaknesses identified")
bullet(d, "Built on synthetic data with no real signal - all correlations ~0, spending uniform across "
          "age/income, K-Means clusters overlapping heavily in PCA space. Most 'insights' are tautological.")
bullet(d, "No cluster validation beyond the elbow (no silhouette, gap statistic, or stability testing).")
bullet(d, "Noise feature (random_numeric_noise) was not explicitly removed before clustering.")
bullet(d, "Deterministic CLV formula is near-circular (CLV almost perfectly tracks Monetary by construction).")
bullet(d, "Raw data contained impossible values (age = -18, negative income) worth scrutinising.")

h1(d, "Conclusion")
para(d, "A solid skeleton and good workflow structure, but a descriptive coding exercise rather than a "
        "statistical study. The upgrade path: real dataset, rigorous cluster validation, method "
        "comparison, and a probabilistic CLV model.")
save(d, "01_Source_Notebook_Analysis.docx")


# ---------------------------------------------------------------------------
# 02 - PROJECT VALUE
# ---------------------------------------------------------------------------
d = new_doc("Project Value", "Why this project helps - business value and career value")
h1(d, "Business value - the core idea")
para(d, "A store cannot afford to treat every customer the same. Without segmentation, marketing is "
        "'spray and pray': the same discount to everyone, wasting money on loyal customers who would "
        "have bought anyway, ignoring customers about to leave, and spending equally on a low-value and "
        "a high-value customer. Segmentation lets a business spend the SAME budget and get MORE back.")

h2(d, "What each step buys the business")
table(d, ["Technical step", "Business use", "Why it makes money"], [
    ["Segmentation (clusters)", "Split customers into groups", "Targeted messaging instead of one wasteful blanket"],
    ["'Champions' segment", "VIP perks, loyalty rewards", "Keep the ~20% of customers driving ~80% of revenue"],
    ["'At Risk' segment", "Win-back offers before churn", "Re-activation is far cheaper than acquisition"],
    ["'Potential' segment", "Nurture with personalised offers", "Largest group = biggest growth opportunity"],
    ["CLV prediction", "Predict future value per customer", "Don't spend more to acquire than a customer is worth"],
])

h2(d, "A concrete example")
para(d, "A GBP 10,000 campaign sent to everyone at 2% response vs the same GBP 10,000 sent only to the "
        "15,000 customers who actually matter at 8% response = roughly 4x the return on the same budget. "
        "That difference is the value. Segmentation tells a company which customers to spend on, how much, "
        "and what to say - instead of guessing.")

h1(d, "Why the statistical rigor matters")
para(d, "Anyone can run K-Means and draw four coloured blobs. The validation, stability testing, and "
        "probabilistic CLV are what make the segments trustworthy enough to bet real money on.")
table(d, ["Without rigor", "With rigor", "Effect"], [
    ["'Here are 4 groups'", "Silhouette + gap statistic confirm 4 is right", "A manager will trust it"],
    ["Might be random slices", "Bootstrap shows the segments reproduce", "Strategy won't fall apart next month"],
    ["Backward-looking formula", "BG/NBD predicts future value", "Plan ahead, not just look back"],
])
para(d, "Rigor is the difference between a pretty chart and a decision a company commits its budget to.")

h1(d, "Career value")
para(d, "Each clause of the project summary is a skill hiring managers screen for, and the project is the proof.")
table(d, ["The work", "Skill it proves"], [
    ["1M+ real transactions cleaned", "Handling messy, real, large data"],
    ["Engineered RFM features", "Turning raw data into business-meaningful variables"],
    ["Validated K (not just elbow)", "Statistical maturity - questioning methods"],
    ["Bootstrap stability", "Reproducibility / robustness"],
    ["K-Means vs GMM vs hierarchical", "Comparing methods and justifying choices"],
    ["Probabilistic BG/NBD CLV", "Advanced statistical modeling"],
    ["Translated to marketing actions", "Connecting statistics to business value"],
])
para(d, "Most segmentation projects on GitHub stop at K-Means + blobs. This one demonstrates the three "
        "things juniors usually can't yet do: validate, compare, and translate.")
save(d, "02_Project_Value.docx")


# ---------------------------------------------------------------------------
# 03 - DATASET
# ---------------------------------------------------------------------------
d = new_doc("The Dataset", "UCI 'Online Retail II' - verified facts")
h1(d, "What it is")
para(d, "Real transaction records from an actual UK-based online retailer selling unique all-occasion "
        "giftware. Notably, many customers are wholesalers, which is why some buy in very large quantities. "
        "One of the most-used real datasets in retail analytics. Source: UCI Machine Learning Repository.")

h1(d, "Verified facts (checked on the downloaded file)")
table(d, ["Property", "Value"], [
    ["Location", "~/customer-segmentation/data/raw/online_retail_II.xlsx (45 MB)"],
    ["Sheets", "Year 2009-2010 (525,461 rows) + Year 2010-2011 (541,910 rows)"],
    ["Total rows", "1,067,371 (one row per product line per invoice)"],
    ["Columns", "8"],
    ["Date span", "01 Dec 2009 -> 09 Dec 2011 (~2 years)"],
    ["Unique customers", "5,942"],
    ["Countries", "43 (heavily UK: 981,330 of 1,067,371 rows)"],
])

h1(d, "The columns")
table(d, ["Column", "Meaning", "Used for"], [
    ["Invoice", "6-digit invoice no; 'C' prefix = cancellation", "Frequency; filter cancellations"],
    ["StockCode", "Product code", "Product-level analysis (optional)"],
    ["Description", "Product name", "Context / sanity checks"],
    ["Quantity", "Units per line; can be negative (returns)", "Monetary; clean negatives"],
    ["InvoiceDate", "Date + time of transaction", "Recency & time trends"],
    ["Price", "Price per unit, GBP", "Monetary"],
    ["Customer ID", "5-digit customer no; many missing", "The unit of analysis (key)"],
    ["Country", "Customer's country", "Geographic segmentation"],
])

h1(d, "Data-quality quirks (the cleaning to-do list)")
table(d, ["Quirk", "Count", "Why it matters"], [
    ["Missing Customer ID", "243,007 (22.8%)", "Can't tie to a customer - dropped for segmentation"],
    ["Cancellations (Invoice 'C')", "19,494", "Not positive purchase behaviour"],
    ["Negative-quantity rows", "22,950", "Returns - clean or net out"],
])
para(d, "These are not bugs - they are realistic messiness that gives genuine, explainable cleaning "
        "decisions in Notebook 01.")

h1(d, "Naming note")
para(d, "'Online Retail' (original) is a single year (~541k rows). 'Online Retail II' is the extended "
        "two-year version (~1M rows) - the one we use, because more history = better CLV predictions.")

h1(d, "Important difference vs the synthetic dataset")
para(d, "The synthetic file had marketing-engagement columns (email_open_rate, click_through_rate, "
        "ad_interactions). This real dataset has none of those - it is pure transaction data. We therefore "
        "lose engagement features but gain real transaction granularity, letting us build RFM ourselves and "
        "run a proper probabilistic CLV model. Clustering features become RFM + tenure + average basket value.")
save(d, "03_Dataset.docx")


# ---------------------------------------------------------------------------
# 04 - COLUMN IMPORTANCE MAP
# ---------------------------------------------------------------------------
d = new_doc("Column Importance Map", "The north star for feature engineering")
h1(d, "The principle")
para(d, "The columns are the problem. Everything downstream (RFM, CLV, clusters) is just a transformation "
        "of these 8 columns. The most common mistake is rushing to KMeans() before knowing what each column "
        "means and how much it deserves to influence the result. Working rule: every feature we create must "
        "trace back to a justified column - no feature without a justification.")

h1(d, "Importance ranking")
table(d, ["Column", "Importance", "What it drives", "If ignored"], [
    ["Customer ID", "5/5", "The unit of analysis", "No segmentation possible at all"],
    ["InvoiceDate", "5/5", "Recency, tenure, BG/NBD timing", "Lose Recency + can't model future value"],
    ["Invoice", "4/5", "Frequency + cancellation flag", "Lose Frequency - half of loyalty gone"],
    ["Quantity x Price", "4/5", "Monetary, avg basket value", "Lose all spend/value signal"],
    ["Country", "2/5", "Geographic context (92% UK)", "Lose a reporting angle, not a core feature"],
    ["StockCode / Description", "1/5", "What each segment buys (optional)", "Lose product insight; RFM unaffected"],
])

h1(d, "The key move: 8 transaction columns -> ~5 customer features")
para(d, "Raw data is one row per product-line. The models need one row per customer. So we collapse many "
        "transaction rows into engineered features:")
table(d, ["Engineered feature", "Built from"], [
    ["Recency", "InvoiceDate (most recent purchase vs snapshot date)"],
    ["Frequency", "Count of distinct Invoice values"],
    ["Monetary", "Sum of Quantity x Price"],
    ["Tenure", "First-to-last InvoiceDate span"],
    ["Average basket value", "Monetary / Frequency"],
])
para(d, "These engineered columns - NOT the raw 8 - are the actual inputs to K-Means and the CLV model. "
        "The project is really about manufacturing the right customer-level columns out of raw transaction "
        "columns. Country and product columns are supporting cast for profiling, not core clustering inputs.")
save(d, "04_Column_Importance_Map.docx")


# ---------------------------------------------------------------------------
# 05 - PROBLEM FORMULATION
# ---------------------------------------------------------------------------
d = new_doc("Problem Formulation", "Formal problem statement, objectives, and expected results")
h1(d, "Background")
para(d, "An online UK gift retailer has two years of transaction history (~1.07M line items, ~5,900 "
        "identifiable customers, 43 countries). It knows individual transactions but has no structured view "
        "of WHO its customers are as groups, which drive revenue, which are leaving, or what each is worth. "
        "Marketing is therefore one-size-fits-all.")

h1(d, "Problem statement")
h2(d, "Business problem")
para(d, "The retailer cannot target marketing effectively because it has no way to distinguish high-value, "
        "loyal, at-risk, and low-value customers, nor to estimate each customer's future worth. Marketing "
        "spend is allocated uniformly rather than where it generates the most return.")
h2(d, "Analytical reframing")
para(d, "Given transaction-level records: (a) construct customer-level behavioral features, (b) discover a "
        "small number of distinct, statistically validated segments using unsupervised clustering, and "
        "(c) model each customer's expected future value probabilistically - then translate both into "
        "actionable marketing strategy. Note: this is unsupervised with NO ground-truth labels, so the "
        "central difficulty is not finding clusters but proving they are real and useful.")

h1(d, "Motivation")
bullet(d, "Pareto reality: ~20% of customers typically drive ~70-80% of revenue.")
bullet(d, "Acquisition vs retention: re-engaging an existing customer is ~5x cheaper than acquiring a new one.")
bullet(d, "Forward-looking gap: historical spend != future value; this requires prediction, not summation.")

h1(d, "Objectives")
table(d, ["#", "Objective", "Met when"], [
    ["O1", "Clean, defensible customer-level dataset", "Documented decisions; reconciled row counts"],
    ["O2", "Engineer behavioral features (RFM + tenure + avg basket)", "One justified row per customer"],
    ["O3", "Discover customer segments via clustering", "A label for every customer"],
    ["O4", "Validate segments are real and stable", "Silhouette, gap statistic, bootstrap reported"],
    ["O5", "Compare clustering methods, justify choice", "K-Means vs GMM vs hierarchical with rationale"],
    ["O6", "Model future customer value probabilistically", "BG/NBD + Gamma-Gamma predicted CLV"],
    ["O7", "Translate findings into marketing actions", "Recommendation table + revenue at stake"],
])

h1(d, "Analytical questions")
numbered(d, "How many genuinely distinct segments exist - and how do we know the number is right?")
numbered(d, "What characterizes each segment (recency, frequency, spend, tenure)?")
numbered(d, "Which segments concentrate the revenue and future value?")
numbered(d, "How much is each customer expected to be worth over the next period?")
numbered(d, "Which customers are at risk of churning, and which are worth investing in?")
numbered(d, "Do data-driven clusters agree with rule-based RFM segments - and where do they disagree?")

h1(d, "Hypotheses (to confirm or refute)")
bullet(d, "H1: A small number of segments (~3-5) optimally describes the customer base.")
bullet(d, "H2: Customer value is highly concentrated - a small 'champion' segment holds a disproportionate "
          "share of revenue and predicted CLV (Pareto-like skew).")
bullet(d, "H3: Frequency and Monetary are positively associated; Recency is largely independent of spend.")
bullet(d, "H4: Density-based clustering (DBSCAN) underperforms centroid/model-based methods because behavior "
          "forms a continuous gradient, not dense islands.")

h1(d, "Scope & assumptions")
bullet(d, "In scope: customers with valid IDs; completed purchases; the two-year window as given.")
bullet(d, "Dropped (with reasons): missing Customer ID (~22.8%); cancellations and pure returns; "
          "zero/negative prices.")
bullet(d, "Assumptions: Recency snapshot = dataset's latest date; CLV horizon = a defined future window (e.g. 12 months).")

h1(d, "Expected results & deliverables")
para(d, "A. Engineered customer table (~5,900 rows): Recency, Frequency, Monetary, Tenure, AvgBasket, plus "
        "Segment, Cluster, PredictedCLV.")
para(d, "B. Segment findings - directional expectation (to be confirmed):")
table(d, ["Likely segment", "Expected profile", "Action"], [
    ["Champions / VIP", "Recent, frequent, high spend (often wholesalers)", "Retain: loyalty perks, early access"],
    ["Loyal / steady", "Regular buyers, moderate spend", "Grow: cross-sell, upsell"],
    ["Potential / new", "Recent but few purchases", "Convert: nurture, onboarding offers"],
    ["At-risk / hibernating", "Bought before, now quiet", "Win back: reminders, discounts"],
])
para(d, "C. Validation results: e.g. 'silhouette and gap statistic both support K=4; bootstrap reproduces "
        "the same segments in >90% of runs' - evidence, not just colored blobs.")
para(d, "D. Method comparison: a table scoring K-Means vs GMM vs hierarchical on validation metrics, with a "
        "justified winner.")
para(d, "E. CLV results: a right-skewed predicted-CLV distribution (most customers low value, a long tail "
        "of high-value ones) and a ranked top-value customer list. High CLV should align with Champions - "
        "an internal consistency check.")
para(d, "F. Final outputs: README, 7 narrated notebooks, src/ package, written report, exported figures, "
        "and the final customer table tying segment + value + action together.")

h1(d, "Success criteria")
para(d, "The project succeeds if it can answer, with evidence: 'How many customer types does this business "
        "have, who are they, what is each worth, and what should marketing do about each?' - and if a reader "
        "can follow why every methodological choice was made. Honest reporting of surprises (e.g. K=3, or "
        "weak clustering) is itself a strong result.")
save(d, "05_Problem_Formulation.docx")


# ---------------------------------------------------------------------------
# 06 - REPO STRUCTURE & ROADMAP
# ---------------------------------------------------------------------------
d = new_doc("Repo Structure & Roadmap", "Planned layout, notebooks, methods, and data flow")
h1(d, "Planned repository layout")
para(d, "customer-segmentation/")
bullet(d, "README.md - storefront: problem, data, methods, key findings, how to run")
bullet(d, "requirements.txt - exact library versions (reproducibility)")
bullet(d, ".gitignore - keep large data / checkpoints out of git")
bullet(d, "data/raw/ - original Online Retail II (never edited - provenance)")
bullet(d, "data/processed/ - cleaned customer table, RFM table")
bullet(d, "docs/ - this planning folder")
bullet(d, "notebooks/ - 01..07, numbered to read in order")
bullet(d, "src/ - reusable modules: data_prep.py, features.py, clustering.py, clv.py")
bullet(d, "reports/figures/ + report - written statistical write-up")
bullet(d, "tests/ - a few unit tests on src/")

h1(d, "The notebooks")
table(d, ["Notebook", "Produces", "What it lets you explain"], [
    ["01 Cleaning", "Clean customer table", "Why cancellations/null IDs dropped; dtype/outlier choices"],
    ["02 EDA", "Figures + insights", "Revenue concentration, country/time patterns, repeat behavior"],
    ["03 RFM", "RFM table + rule segments", "R/F/M meaning, inverted recency, quintile binning"],
    ["04 Clustering", "K-Means + GMM + hierarchical fits", "Why scale first, what PCA shows"],
    ["05 Validation", "Silhouette / gap / stability", "Proof K is right and clusters are stable"],
    ["06 CLV", "BG/NBD + Gamma-Gamma predictions", "Probabilistic prediction of future value"],
    ["07 Segments", "Named segments + recommendations", "Who each segment is, what to do, revenue at stake"],
])

h1(d, "Statistical methods map")
table(d, ["Method", "Why used", "What it demonstrates"], [
    ["Median imputation, IQR capping", "Robust to skew/outliers", "Knowledge of distributional assumptions"],
    ["RFM", "Standard behavioral segmentation", "Domain feature engineering"],
    ["StandardScaler", "K-Means is distance-based", "Understanding the algorithm's needs"],
    ["Silhouette + gap statistic", "Choose K with evidence", "Model-selection rigor"],
    ["Bootstrap stability", "Are clusters reproducible?", "Testing, not assuming"],
    ["K-Means vs GMM vs hierarchical", "Different assumptions", "Method comparison & justification"],
    ["BG/NBD + Gamma-Gamma", "Probabilistic vs deterministic CLV", "Advanced statistical modeling"],
    ["PCA", "2-D visualization of clusters", "Dimensionality reduction"],
])

h1(d, "End-to-end data flow")
para(d, "raw transactions (1M+ rows)")
para(d, "  -> 01: clean -> customer-transaction table")
para(d, "    -> 03: aggregate -> RFM table (one row per customer)")
para(d, "      -> 04+05: cluster & validate -> segment label per customer")
para(d, "      -> 06: BG/NBD model -> predicted CLV per customer")
para(d, "        -> 07: join -> final customer table (segment + CLV + action)")
para(d, "Every customer ends with: RFM scores, a validated cluster label, a predicted lifetime value, "
        "and a recommended marketing action. That final table is the project's payoff.")

h1(d, "Working style")
bullet(d, "Build phase by phase; review and understand each before moving on.")
bullet(d, "Every methodological choice must be explainable.")
bullet(d, "Next concrete step: Notebook 01 (cleaning).")
save(d, "06_Repo_Structure_and_Roadmap.docx")

print("\nAll documents generated in:", OUT_DIR)

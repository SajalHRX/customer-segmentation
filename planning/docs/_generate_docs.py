"""
Generates the Customer Segmentation project planning documents as Word .docx files.
Each document captures one topic from our planning conversation.
Run:  python3 _generate_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # dark blue
GREY = RGBColor(0x55, 0x55, 0x55)


def new_doc(title, subtitle):
    doc = Document()
    # base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_heading(title, level=0)
    for run in t.runs:
        run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    r = sub.add_run(subtitle)
    r.italic = True
    r.font.color.rgb = GREY
    r.font.size = Pt(11)

    meta = doc.add_paragraph()
    m = meta.add_run("Customer Segmentation Project  |  Planning Notes  |  17 June 2026")
    m.font.size = Pt(9)
    m.font.color.rgb = GREY
    doc.add_paragraph("")
    return doc


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def para(doc, text):
    return doc.add_paragraph(text)


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph("")
    return t


def save(doc, name):
    path = os.path.join(OUT_DIR, name)
    doc.save(path)
    print("wrote", name)


# ---------------------------------------------------------------------------
# 00 - PROJECT OVERVIEW (index)
# ---------------------------------------------------------------------------
d = new_doc("Project Overview", "What we are building and why")
h1(d, "One-line summary")
para(d, "A rigorous, explainable customer-segmentation project built on real e-commerce "
        "transaction data (UCI 'Online Retail II'). It segments customers into a small number "
        "of statistically validated groups, predicts each customer's future value with a "
        "probabilistic model, and translates the result into concrete marketing actions.")

h1(d, "The interview-ready story")
para(d, "\"I took 1M+ real e-commerce transactions, cleaned them into customer-level records, "
        "engineered RFM features, and segmented customers using K-Means - but instead of trusting "
        "the elbow method blindly, I validated the number of clusters with silhouette and gap "
        "statistics, tested cluster stability with bootstrapping, and benchmarked K-Means against "
        "Gaussian Mixture and hierarchical clustering. I then replaced the naive lifetime-value "
        "formula with a probabilistic BG/NBD + Gamma-Gamma model to predict each customer's future "
        "value, and translated the segments into concrete marketing actions.\"")

h1(d, "Project decisions (locked in)")
table(d, ["Decision", "Choice"], [
    ["Goal", "Portfolio / CV piece"],
    ["Dataset", "Real - UCI 'Online Retail II' (downloaded & verified)"],
    ["Statistical rigor", "Full - validation, method comparison, probabilistic CLV"],
    ["Format of work", "GitHub repo, numbered notebooks, reusable src/ modules"],
    ["Working style", "Phase by phase; everything must be explainable"],
])

h1(d, "Document index")
para(d, "This folder contains the planning conversation, split by topic:")
table(d, ["Document", "Contents"], [
    ["00_Project_Overview", "This file - summary, decisions, index"],
    ["01_Source_Notebook_Analysis", "Analysis of the original synthetic-data notebook and its weaknesses"],
    ["02_Project_Value", "Business value + career value; why statistical rigor matters"],
    ["03_Dataset", "Online Retail II - verified facts, columns, data-quality quirks"],
    ["04_Column_Importance_Map", "Role and importance of every column; the feature-engineering principle"],
    ["05_Problem_Formulation", "Formal problem statement, objectives, hypotheses, expected results"],
    ["06_Repo_Structure_and_Roadmap", "Planned repo layout, notebooks, methods map, data flow"],
    ["07_Problem_Framing_and_Scope", "Design discussion #1 - chosen framing (retention/budget), honest scope boundaries"],
    ["08_Time_Window_and_CLV_Validation", "Design discussion #2 - anchor date, temporal holdout, multi-cutoff experiment"],
    ["09_Feature_Treatment", "Design discussion #3 - skew/log, scaling (Robust vs Standard), weighting & redundancy"],
    ["10_Choosing_and_Validating_K", "Design discussion #4 - choosing K (elbow/silhouette/gap/CH/DB) + stability (bootstrap/consensus)"],
    ["11_Clustering_Method_Comparison", "Design discussion #5 - K-Means vs GMM vs hierarchical; the continuum question; ARI/BIC"],
    ["12_CLV_and_Segments_Integration", "Design discussion #6 - CLV as a separate post-hoc layer joined on Customer ID; the segment x CLV grid"],
    ["13_Segment_Profiling_and_Validation", "Design discussion #7a - profiling clusters into named personas; validating they are real (3 tiers)"],
    ["14_Recommendations_and_Success_Criteria", "Design discussion #7b - the action/value-at-stake grid; what 'good' means for the project"],
])
para(d, "Documents 07+ are DESIGN DISCUSSIONS - decisions worked out topic-by-topic after planning, "
        "each recording the reasoning behind a methodological choice.")
save(d, "00_Project_Overview.docx")


# ---------------------------------------------------------------------------
# 01 - SOURCE NOTEBOOK ANALYSIS
# ---------------------------------------------------------------------------
d = new_doc("Source Notebook Analysis", "Understanding the original synthetic-data notebook")
h1(d, "What the original notebook was")
para(d, "The project started from an existing 58-page notebook titled "
        "'Marketing_Analytics_&_Customer_Segmentation'. It ran an end-to-end pipeline in Python "
        "(pandas, numpy, matplotlib, seaborn, scikit-learn) on a SYNTHETIC dataset of 125,000 "
        "customers x 14 columns (marketing_customer.csv).")

h1(d, "The pipeline it followed")
numbered(d, "Cleaning: impute missing categoricals with 'unknown', numerics with median; fix dtypes; "
            "IQR outlier capping (winsorization); drop 5,000 duplicate rows; drop a junk text column.")
numbered(d, "EDA: histograms, boxplots, categorical counts, scatterplots, revenue by channel/location, "
            "correlation heatmap (all correlations near zero - by design).")
numbered(d, "Feature engineering: RFM (Recency, Frequency, Monetary), 1-5 RFM scores, rule-based "
            "segments (Champions/Loyal/Potential/At-Risk/Lost), and a deterministic CLV formula.")
numbered(d, "Segmentation: K-Means (elbow -> K=4, scaled features, PCA viz) and DBSCAN (collapsed to "
            "one blob - correctly judged unsuitable).")
numbered(d, "Insights & recommendations: marketing actions per segment.")

h1(d, "Key weaknesses identified")
bullet(d, "Built on synthetic data with no real signal - all correlations ~0, spending uniform across "
          "age/income, K-Means clusters overlapping heavily in PCA space. Most 'insights' are tautological.")
bullet(d, "No cluster validation beyond the elbow (no silhouette, gap statistic, or stability testing).")
bullet(d, "Noise feature (random_numeric_noise) was not explicitly removed before clustering.")
bullet(d, "Deterministic CLV formula is near-circular (CLV almost perfectly tracks Monetary by construction).")
bullet(d, "Raw data contained impossible values (age = -18, negative income) worth scrutinising.")

h1(d, "Conclusion")
para(d, "A solid skeleton and good workflow structure, but a descriptive coding exercise rather than a "
        "statistical study. The upgrade path: real dataset, rigorous cluster validation, method "
        "comparison, and a probabilistic CLV model.")
save(d, "01_Source_Notebook_Analysis.docx")


# ---------------------------------------------------------------------------
# 02 - PROJECT VALUE
# ---------------------------------------------------------------------------
d = new_doc("Project Value", "Why this project helps - business value and career value")
h1(d, "Business value - the core idea")
para(d, "A store cannot afford to treat every customer the same. Without segmentation, marketing is "
        "'spray and pray': the same discount to everyone, wasting money on loyal customers who would "
        "have bought anyway, ignoring customers about to leave, and spending equally on a low-value and "
        "a high-value customer. Segmentation lets a business spend the SAME budget and get MORE back.")

h2(d, "What each step buys the business")
table(d, ["Technical step", "Business use", "Why it makes money"], [
    ["Segmentation (clusters)", "Split customers into groups", "Targeted messaging instead of one wasteful blanket"],
    ["'Champions' segment", "VIP perks, loyalty rewards", "Keep the ~20% of customers driving ~80% of revenue"],
    ["'At Risk' segment", "Win-back offers before churn", "Re-activation is far cheaper than acquisition"],
    ["'Potential' segment", "Nurture with personalised offers", "Largest group = biggest growth opportunity"],
    ["CLV prediction", "Predict future value per customer", "Don't spend more to acquire than a customer is worth"],
])

h2(d, "A concrete example")
para(d, "A GBP 10,000 campaign sent to everyone at 2% response vs the same GBP 10,000 sent only to the "
        "15,000 customers who actually matter at 8% response = roughly 4x the return on the same budget. "
        "That difference is the value. Segmentation tells a company which customers to spend on, how much, "
        "and what to say - instead of guessing.")

h1(d, "Why the statistical rigor matters")
para(d, "Anyone can run K-Means and draw four coloured blobs. The validation, stability testing, and "
        "probabilistic CLV are what make the segments trustworthy enough to bet real money on.")
table(d, ["Without rigor", "With rigor", "Effect"], [
    ["'Here are 4 groups'", "Silhouette + gap statistic confirm 4 is right", "A manager will trust it"],
    ["Might be random slices", "Bootstrap shows the segments reproduce", "Strategy won't fall apart next month"],
    ["Backward-looking formula", "BG/NBD predicts future value", "Plan ahead, not just look back"],
])
para(d, "Rigor is the difference between a pretty chart and a decision a company commits its budget to.")

h1(d, "Career value")
para(d, "Each clause of the project summary is a skill hiring managers screen for, and the project is the proof.")
table(d, ["The work", "Skill it proves"], [
    ["1M+ real transactions cleaned", "Handling messy, real, large data"],
    ["Engineered RFM features", "Turning raw data into business-meaningful variables"],
    ["Validated K (not just elbow)", "Statistical maturity - questioning methods"],
    ["Bootstrap stability", "Reproducibility / robustness"],
    ["K-Means vs GMM vs hierarchical", "Comparing methods and justifying choices"],
    ["Probabilistic BG/NBD CLV", "Advanced statistical modeling"],
    ["Translated to marketing actions", "Connecting statistics to business value"],
])
para(d, "Most segmentation projects on GitHub stop at K-Means + blobs. This one demonstrates the three "
        "things juniors usually can't yet do: validate, compare, and translate.")
save(d, "02_Project_Value.docx")


# ---------------------------------------------------------------------------
# 03 - DATASET
# ---------------------------------------------------------------------------
d = new_doc("The Dataset", "UCI 'Online Retail II' - verified facts")
h1(d, "What it is")
para(d, "Real transaction records from an actual UK-based online retailer selling unique all-occasion "
        "giftware. Notably, many customers are wholesalers, which is why some buy in very large quantities. "
        "One of the most-used real datasets in retail analytics. Source: UCI Machine Learning Repository.")

h1(d, "Verified facts (checked on the downloaded file)")
table(d, ["Property", "Value"], [
    ["Location", "~/customer-segmentation/data/raw/online_retail_II.xlsx (45 MB)"],
    ["Sheets", "Year 2009-2010 (525,461 rows) + Year 2010-2011 (541,910 rows)"],
    ["Total rows", "1,067,371 (one row per product line per invoice)"],
    ["Columns", "8"],
    ["Date span", "01 Dec 2009 -> 09 Dec 2011 (~2 years)"],
    ["Unique customers", "5,942"],
    ["Countries", "43 (heavily UK: 981,330 of 1,067,371 rows)"],
])

h1(d, "The columns")
table(d, ["Column", "Meaning", "Used for"], [
    ["Invoice", "6-digit invoice no; 'C' prefix = cancellation", "Frequency; filter cancellations"],
    ["StockCode", "Product code", "Product-level analysis (optional)"],
    ["Description", "Product name", "Context / sanity checks"],
    ["Quantity", "Units per line; can be negative (returns)", "Monetary; clean negatives"],
    ["InvoiceDate", "Date + time of transaction", "Recency & time trends"],
    ["Price", "Price per unit, GBP", "Monetary"],
    ["Customer ID", "5-digit customer no; many missing", "The unit of analysis (key)"],
    ["Country", "Customer's country", "Geographic segmentation"],
])

h1(d, "Data-quality quirks (the cleaning to-do list)")
table(d, ["Quirk", "Count", "Why it matters"], [
    ["Missing Customer ID", "243,007 (22.8%)", "Can't tie to a customer - dropped for segmentation"],
    ["Cancellations (Invoice 'C')", "19,494", "Not positive purchase behaviour"],
    ["Negative-quantity rows", "22,950", "Returns - clean or net out"],
])
para(d, "These are not bugs - they are realistic messiness that gives genuine, explainable cleaning "
        "decisions in Notebook 01.")

h1(d, "Naming note")
para(d, "'Online Retail' (original) is a single year (~541k rows). 'Online Retail II' is the extended "
        "two-year version (~1M rows) - the one we use, because more history = better CLV predictions.")

h1(d, "Important difference vs the synthetic dataset")
para(d, "The synthetic file had marketing-engagement columns (email_open_rate, click_through_rate, "
        "ad_interactions). This real dataset has none of those - it is pure transaction data. We therefore "
        "lose engagement features but gain real transaction granularity, letting us build RFM ourselves and "
        "run a proper probabilistic CLV model. Clustering features become RFM + tenure + average basket value.")
save(d, "03_Dataset.docx")


# ---------------------------------------------------------------------------
# 04 - COLUMN IMPORTANCE MAP
# ---------------------------------------------------------------------------
d = new_doc("Column Importance Map", "The north star for feature engineering")
h1(d, "The principle")
para(d, "The columns are the problem. Everything downstream (RFM, CLV, clusters) is just a transformation "
        "of these 8 columns. The most common mistake is rushing to KMeans() before knowing what each column "
        "means and how much it deserves to influence the result. Working rule: every feature we create must "
        "trace back to a justified column - no feature without a justification.")

h1(d, "Importance ranking")
table(d, ["Column", "Importance", "What it drives", "If ignored"], [
    ["Customer ID", "5/5", "The unit of analysis", "No segmentation possible at all"],
    ["InvoiceDate", "5/5", "Recency, tenure, BG/NBD timing", "Lose Recency + can't model future value"],
    ["Invoice", "4/5", "Frequency + cancellation flag", "Lose Frequency - half of loyalty gone"],
    ["Quantity x Price", "4/5", "Monetary, avg basket value", "Lose all spend/value signal"],
    ["Country", "2/5", "Geographic context (92% UK)", "Lose a reporting angle, not a core feature"],
    ["StockCode / Description", "1/5", "What each segment buys (optional)", "Lose product insight; RFM unaffected"],
])

h1(d, "The key move: 8 transaction columns -> ~5 customer features")
para(d, "Raw data is one row per product-line. The models need one row per customer. So we collapse many "
        "transaction rows into engineered features:")
table(d, ["Engineered feature", "Built from"], [
    ["Recency", "InvoiceDate (most recent purchase vs snapshot date)"],
    ["Frequency", "Count of distinct Invoice values"],
    ["Monetary", "Sum of Quantity x Price"],
    ["Tenure", "First-to-last InvoiceDate span"],
    ["Average basket value", "Monetary / Frequency"],
])
para(d, "These engineered columns - NOT the raw 8 - are the actual inputs to K-Means and the CLV model. "
        "The project is really about manufacturing the right customer-level columns out of raw transaction "
        "columns. Country and product columns are supporting cast for profiling, not core clustering inputs.")
save(d, "04_Column_Importance_Map.docx")


# ---------------------------------------------------------------------------
# 05 - PROBLEM FORMULATION
# ---------------------------------------------------------------------------
d = new_doc("Problem Formulation", "Formal problem statement, objectives, and expected results")
h1(d, "Background")
para(d, "An online UK gift retailer has two years of transaction history (~1.07M line items, ~5,900 "
        "identifiable customers, 43 countries). It knows individual transactions but has no structured view "
        "of WHO its customers are as groups, which drive revenue, which are leaving, or what each is worth. "
        "Marketing is therefore one-size-fits-all.")

h1(d, "Problem statement")
h2(d, "Business problem")
para(d, "The retailer cannot target marketing effectively because it has no way to distinguish high-value, "
        "loyal, at-risk, and low-value customers, nor to estimate each customer's future worth. Marketing "
        "spend is allocated uniformly rather than where it generates the most return.")
h2(d, "Analytical reframing")
para(d, "Given transaction-level records: (a) construct customer-level behavioral features, (b) discover a "
        "small number of distinct, statistically validated segments using unsupervised clustering, and "
        "(c) model each customer's expected future value probabilistically - then translate both into "
        "actionable marketing strategy. Note: this is unsupervised with NO ground-truth labels, so the "
        "central difficulty is not finding clusters but proving they are real and useful.")

h1(d, "Motivation")
bullet(d, "Pareto reality: ~20% of customers typically drive ~70-80% of revenue.")
bullet(d, "Acquisition vs retention: re-engaging an existing customer is ~5x cheaper than acquiring a new one.")
bullet(d, "Forward-looking gap: historical spend != future value; this requires prediction, not summation.")

h1(d, "Objectives")
table(d, ["#", "Objective", "Met when"], [
    ["O1", "Clean, defensible customer-level dataset", "Documented decisions; reconciled row counts"],
    ["O2", "Engineer behavioral features (RFM + tenure + avg basket)", "One justified row per customer"],
    ["O3", "Discover customer segments via clustering", "A label for every customer"],
    ["O4", "Validate segments are real and stable", "Silhouette, gap statistic, bootstrap reported"],
    ["O5", "Compare clustering methods, justify choice", "K-Means vs GMM vs hierarchical with rationale"],
    ["O6", "Model future customer value probabilistically", "BG/NBD + Gamma-Gamma predicted CLV"],
    ["O7", "Translate findings into marketing actions", "Recommendation table + revenue at stake"],
])

h1(d, "Analytical questions")
numbered(d, "How many genuinely distinct segments exist - and how do we know the number is right?")
numbered(d, "What characterizes each segment (recency, frequency, spend, tenure)?")
numbered(d, "Which segments concentrate the revenue and future value?")
numbered(d, "How much is each customer expected to be worth over the next period?")
numbered(d, "Which customers are at risk of churning, and which are worth investing in?")
numbered(d, "Do data-driven clusters agree with rule-based RFM segments - and where do they disagree?")

h1(d, "Hypotheses (to confirm or refute)")
bullet(d, "H1: A small number of segments (~3-5) optimally describes the customer base.")
bullet(d, "H2: Customer value is highly concentrated - a small 'champion' segment holds a disproportionate "
          "share of revenue and predicted CLV (Pareto-like skew).")
bullet(d, "H3: Frequency and Monetary are positively associated; Recency is largely independent of spend.")
bullet(d, "H4: Density-based clustering (DBSCAN) underperforms centroid/model-based methods because behavior "
          "forms a continuous gradient, not dense islands.")

h1(d, "Scope & assumptions")
bullet(d, "In scope: customers with valid IDs; completed purchases; the two-year window as given.")
bullet(d, "Dropped (with reasons): missing Customer ID (~22.8%); cancellations and pure returns; "
          "zero/negative prices.")
bullet(d, "Assumptions: Recency snapshot = dataset's latest date; CLV horizon = a defined future window (e.g. 12 months).")

h1(d, "Expected results & deliverables")
para(d, "A. Engineered customer table (~5,900 rows): Recency, Frequency, Monetary, Tenure, AvgBasket, plus "
        "Segment, Cluster, PredictedCLV.")
para(d, "B. Segment findings - directional expectation (to be confirmed):")
table(d, ["Likely segment", "Expected profile", "Action"], [
    ["Champions / VIP", "Recent, frequent, high spend (often wholesalers)", "Retain: loyalty perks, early access"],
    ["Loyal / steady", "Regular buyers, moderate spend", "Grow: cross-sell, upsell"],
    ["Potential / new", "Recent but few purchases", "Convert: nurture, onboarding offers"],
    ["At-risk / hibernating", "Bought before, now quiet", "Win back: reminders, discounts"],
])
para(d, "C. Validation results: e.g. 'silhouette and gap statistic both support K=4; bootstrap reproduces "
        "the same segments in >90% of runs' - evidence, not just colored blobs.")
para(d, "D. Method comparison: a table scoring K-Means vs GMM vs hierarchical on validation metrics, with a "
        "justified winner.")
para(d, "E. CLV results: a right-skewed predicted-CLV distribution (most customers low value, a long tail "
        "of high-value ones) and a ranked top-value customer list. High CLV should align with Champions - "
        "an internal consistency check.")
para(d, "F. Final outputs: README, 7 narrated notebooks, src/ package, written report, exported figures, "
        "and the final customer table tying segment + value + action together.")

h1(d, "Success criteria")
para(d, "The project succeeds if it can answer, with evidence: 'How many customer types does this business "
        "have, who are they, what is each worth, and what should marketing do about each?' - and if a reader "
        "can follow why every methodological choice was made. Honest reporting of surprises (e.g. K=3, or "
        "weak clustering) is itself a strong result.")
save(d, "05_Problem_Formulation.docx")


# ---------------------------------------------------------------------------
# 06 - REPO STRUCTURE & ROADMAP
# ---------------------------------------------------------------------------
d = new_doc("Repo Structure & Roadmap", "Planned layout, notebooks, methods, and data flow")
h1(d, "Planned repository layout")
para(d, "customer-segmentation/")
bullet(d, "README.md - storefront: problem, data, methods, key findings, how to run")
bullet(d, "requirements.txt - exact library versions (reproducibility)")
bullet(d, ".gitignore - keep large data / checkpoints out of git")
bullet(d, "data/raw/ - original Online Retail II (never edited - provenance)")
bullet(d, "data/processed/ - cleaned customer table, RFM table")
bullet(d, "docs/ - this planning folder")
bullet(d, "notebooks/ - 01..07, numbered to read in order")
bullet(d, "src/ - reusable modules: data_prep.py, features.py, clustering.py, clv.py")
bullet(d, "reports/figures/ + report - written statistical write-up")
bullet(d, "tests/ - a few unit tests on src/")

h1(d, "The notebooks")
table(d, ["Notebook", "Produces", "What it lets you explain"], [
    ["01 Cleaning", "Clean customer table", "Why cancellations/null IDs dropped; dtype/outlier choices"],
    ["02 EDA", "Figures + insights", "Revenue concentration, country/time patterns, repeat behavior"],
    ["03 RFM", "RFM table + rule segments", "R/F/M meaning, inverted recency, quintile binning"],
    ["04 Clustering", "K-Means + GMM + hierarchical fits", "Why scale first, what PCA shows"],
    ["05 Validation", "Silhouette / gap / stability", "Proof K is right and clusters are stable"],
    ["06 CLV", "BG/NBD + Gamma-Gamma predictions", "Probabilistic prediction of future value"],
    ["07 Segments", "Named segments + recommendations", "Who each segment is, what to do, revenue at stake"],
])

h1(d, "Statistical methods map")
table(d, ["Method", "Why used", "What it demonstrates"], [
    ["Median imputation, IQR capping", "Robust to skew/outliers", "Knowledge of distributional assumptions"],
    ["RFM", "Standard behavioral segmentation", "Domain feature engineering"],
    ["StandardScaler", "K-Means is distance-based", "Understanding the algorithm's needs"],
    ["Silhouette + gap statistic", "Choose K with evidence", "Model-selection rigor"],
    ["Bootstrap stability", "Are clusters reproducible?", "Testing, not assuming"],
    ["K-Means vs GMM vs hierarchical", "Different assumptions", "Method comparison & justification"],
    ["BG/NBD + Gamma-Gamma", "Probabilistic vs deterministic CLV", "Advanced statistical modeling"],
    ["PCA", "2-D visualization of clusters", "Dimensionality reduction"],
])

h1(d, "End-to-end data flow")
para(d, "raw transactions (1M+ rows)")
para(d, "  -> 01: clean -> customer-transaction table")
para(d, "    -> 03: aggregate -> RFM table (one row per customer)")
para(d, "      -> 04+05: cluster & validate -> segment label per customer")
para(d, "      -> 06: BG/NBD model -> predicted CLV per customer")
para(d, "        -> 07: join -> final customer table (segment + CLV + action)")
para(d, "Every customer ends with: RFM scores, a validated cluster label, a predicted lifetime value, "
        "and a recommended marketing action. That final table is the project's payoff.")

h1(d, "Working style")
bullet(d, "Build phase by phase; review and understand each before moving on.")
bullet(d, "Every methodological choice must be explainable.")
bullet(d, "Next concrete step: Notebook 01 (cleaning).")
save(d, "06_Repo_Structure_and_Roadmap.docx")


# ---------------------------------------------------------------------------
# 07 - PROBLEM FRAMING & SCOPE  (Design discussion #1, 18 June 2026)
# ---------------------------------------------------------------------------
d = new_doc("Problem Framing & Scope", "Design discussion #1 - what decision the segmentation drives")
para(d, "Added 18 June 2026. This doc records the first design discussion: pinning down WHO uses the "
        "segmentation and WHAT decision it resolves, before writing any code.")

h1(d, "The question we had to answer first")
para(d, "\"Who runs this business, and what decision does the segmentation help them make?\" The data is a "
        "UK-based online gift/homeware retailer (2009-2011), wholesale + retail mix, mostly UK, many repeat "
        "buyers. That context constrains which framings are HONEST - we must not invent signals the data "
        "does not contain (no ad clicks, no web sessions, pure transactions).")

h1(d, "Three framings considered")
table(d, ["Framing", "Idea", "Verdict"], [
    ["A. Retention & budget allocation", "Group by behaviour so a fixed marketing budget goes where it pays off",
     "CHOSEN - every planned technique (RFM, validated clustering, CLV) has a clear job"],
    ["B. VIP / high-value identification", "Find the ~20% driving ~80% of revenue; build a loyalty tier",
     "Too narrow - risks collapsing to 'rank by Monetary', undersells clustering"],
    ["C. Product / market expansion", "Segment to find underserved customer types or geographies",
     "Weakest - geography is 92% UK; product columns deliberately down-weighted"],
])

h1(d, "Chosen framing (A) - locked problem statement")
para(d, "A UK-based online retailer has a finite marketing/CRM budget and currently treats all customers "
        "alike. Using only transactional history, segment customers by purchasing behaviour (RFM + tenure + "
        "basket value), validate the segments are real and stable, estimate each customer's future value "
        "probabilistically (CLV), and recommend a differentiated action per segment - so retention/marketing "
        "spend follows expected return.")

h1(d, "Stress-test: 'If I were Ajio, would this resolve a real problem?'")
para(d, "Yes - but a SPECIFIC problem, not the whole business. Behavioural (RFM) segmentation is genuinely "
        "deployed in production CRM (Klaviyo, Salesforce Marketing Cloud, etc.). It resolves: 'Given a finite "
        "budget, which customers get which treatment?'")
h2(d, "What framing A RESOLVES (per segment, the decision it answers)")
table(d, ["Segment (behaviour)", "Decision resolved"], [
    ["Champions (recent, frequent, high spend)", "Don't waste discounts - early access, loyalty perks. Protect."],
    ["At-risk loyalists (were active, gone quiet)", "Send win-back offer now - worth saving. Intervene."],
    ["New customers (1 recent order)", "Nurture into a second purchase. Onboard."],
    ["Bargain one-timers (low value, churned)", "Don't chase. Deprioritise."],
    ["Big spenders, infrequent", "Re-engage at the right cadence. Reactivate."],
])
h2(d, "What framing A does NOT solve (state this honestly in interviews)")
bullet(d, "WHAT product to recommend to a user -> that is a recommendation system (collaborative filtering).")
bullet(d, "WHY a customer churned -> needs richer signals (browsing, returns, complaints).")
bullet(d, "WHICH channel / creative converts -> needs experimentation (A/B tests).")
para(d, "It answers 'who, and how much effort' - not 'what exactly to show them'. Those are complementary "
        "layers of the personalisation stack.")

h1(d, "Why this is the right scope for our data")
para(d, "Our data is purely transactional (invoices, quantities, prices, dates) - no clicks, no product "
        "views, no ad data. So we CAN do behavioural value/lifecycle segmentation rigorously (framing A); we "
        "CANNOT honestly build a recommender or churn-cause model without inventing signals. Framing A is "
        "therefore the most ambitious HONEST framing the data supports - and saying out loud what the data "
        "can and cannot answer is itself a mark of analytical judgement.")
save(d, "07_Problem_Framing_and_Scope.docx")


# ---------------------------------------------------------------------------
# 08 - TIME WINDOW & CLV VALIDATION  (Design discussion #2, 18 June 2026)
# ---------------------------------------------------------------------------
d = new_doc("Time Window & CLV Validation", "Design discussion #2 - the anchor date and how we validate CLV")
para(d, "Added 18 June 2026. Records two linked decisions: what 'today' means for Recency/CLV, and how the "
        "probabilistic CLV model is validated out-of-sample.")

h1(d, "The anchor date ('what is today?')")
para(d, "Recency = days since the customer's last purchase - but days since WHEN? The real today (2026) is "
        "useless: the data ends 09 Dec 2011, so every customer would look ~14 years lapsed. The anchor must "
        "sit ON OR AFTER the last transaction, or some customers get negative recency.")
table(d, ["Option", "Verdict"], [
    ["Real today (2026)", "Rejected - no recency variation that matters"],
    ["Last txn date (2011-12-10)", "Valid, but slightly arbitrary mid-month point"],
    ["End-of-FY March 2012", "Rejected - no Jan-Mar 2012 data; adds 3 empty months of fake lapse"],
    ["01 Jan 2012  (CHOSEN)", "Clean 'turn-of-year snapshot'; satisfies >= last txn; uniform ~3-week offset"],
])
para(d, "Decision: anchor = 01 Jan 2012. Narrative: 'Standing at the start of 2012, looking back over all "
        "history.' The uniform offset does not distort the RELATIVE recency ordering that clustering uses.")

h1(d, "Two DIFFERENT validations (do not confuse them)")
table(d, ["Validation", "Validates", "Method"], [
    ["Cluster validation (later step)", "Are the SEGMENTS real and stable?", "silhouette, gap statistic, bootstrap"],
    ["CLV model validation (this doc)", "Does the PREDICTIVE model actually predict?", "temporal holdout"],
])

h1(d, "CLV validation = temporal holdout (train on the past, test on the future)")
para(d, "The point of BG/NBD + Gamma-Gamma is to predict future behaviour, so we hide the future and check "
        "the model's guess.")
numbered(d, "Split the timeline at a cutoff into CALIBRATION (fit) and HOLDOUT (test) periods.")
numbered(d, "Fit BG/NBD using ONLY the calibration period (frequency, recency, tenure as of the cutoff).")
numbered(d, "Predict each customer's expected number of purchases in the holdout window.")
numbered(d, "Compare predicted vs actual holdout behaviour. (lifetimes.calibration_and_holdout_data() does the split.)")
h2(d, "Metrics recorded")
bullet(d, "Calibration plot: predicted vs actual per frequency-bin - a good model hugs the diagonal (main visual).")
bullet(d, "Aggregate: total predicted vs total actual transactions in holdout.")
bullet(d, "Per-customer error: MAE / RMSE, normalised per month of holdout so cutoffs are comparable.")
bullet(d, "Fitted BG/NBD parameters (r, alpha, a, b): do they stay stable across splits? Big swings = fragile.")
bullet(d, "Gamma-Gamma: predicted vs actual avg order value; AND check its key assumption that frequency and "
          "monetary value are ~uncorrelated (compute the correlation; should be near zero).")

h1(d, "Experiment: multiple cutoffs (holdouts of 3, 6, 9 months)")
para(d, "Rather than one holdout, run three to turn validation into a sensitivity analysis (and for practice).")
table(d, ["Holdout", "Cutoff (approx)", "Calibration", "Prediction horizon"], [
    ["3 months", "~2011-09-09", "~21 months (more training)", "short -> easier to predict"],
    ["6 months", "~2011-06-09", "~18 months", "medium"],
    ["9 months", "~2011-03-09", "~15 months (less training)", "long -> harder to predict"],
])
h2(d, "The trap to avoid (the key insight)")
para(d, "Moving the cutoff changes TWO things at once: calibration length AND prediction horizon. The "
        "3-month model will almost certainly show lower raw error - not because it is a better model, but "
        "because predicting 3 months ahead is easier and it had more training data. So do NOT rank models by "
        "raw MAE and declare '3-month is best' - that is the wrong conclusion.")
h2(d, "How to read it honestly")
bullet(d, "Treat each cutoff as a different question ('how well does it predict 3 / 6 / 9 months out?'); the "
          "interesting result is how GRACEFULLY accuracy decays with horizon, not which wins.")
bullet(d, "Judge CALIBRATION (predicted vs actual by bin) and PARAMETER STABILITY, which are far more "
          "horizon-fair than raw error. Diagonal-hugging across all three cutoffs = strong robustness evidence.")
bullet(d, "Caveat: the customer cohort shifts slightly per cutoff (anyone whose first purchase is after the "
          "cutoff is excluded from calibration). Note this in the writeup.")
h2(d, "Engineering note")
para(d, "Build a reusable function evaluate_cutoff(months) -> {metrics, plot}, run over [3, 6, 9], collected "
        "into one comparison table + a row of calibration small-multiples. Cleaner than copy-paste and "
        "demonstrates a reproducible experiment design.")

h1(d, "How validation fits with the production model")
bullet(d, "Segmentation: fit on FULL history (all data up to 2012-01-01) for maximum signal.")
bullet(d, "CLV validation: a SEPARATE holdout experiment to prove the model is trustworthy.")
bullet(d, "Final CLV numbers attached to customers/segments: refit on full history once validation earns trust.")
para(d, "So the holdout is a credibility experiment; the production CLV uses all the data. 'The model stays "
        "calibrated from 3 to 9 months and its parameters are stable' is a much stronger claim than a single "
        "holdout.")
save(d, "08_Time_Window_and_CLV_Validation.docx")


# ---------------------------------------------------------------------------
# 09 - FEATURE TREATMENT  (Design discussion #3, 18 June 2026)
# ---------------------------------------------------------------------------
d = new_doc("Feature Treatment", "Design discussion #3 - skew, scaling, and feature weighting")
para(d, "Added 18 June 2026. How the customer-level features are transformed before clustering. Three "
        "linked issues: distributional skew, scaling, and weighting/redundancy. Standing rule: every "
        "experiment below must be observable (compare variants, never run blind).")

h1(d, "Issue 1 - Skew (the log transform)")
para(d, "RFM features are strongly right-skewed (Pareto-like): most customers buy a little, a few "
        "wholesalers buy enormously. Monetary can span ~5 orders of magnitude (~GBP 3 to ~GBP 280,000).")
h2(d, "Why it breaks K-Means")
bullet(d, "Euclidean distance sums squared differences; a feature with huge spread dominates the distance "
          "(raw Monetary swamps Recency) - clustering silently collapses to 'sort by spend'.")
bullet(d, "Even after scaling, a few whales sit so far out that K-Means wastes whole clusters isolating "
          "them: {whale1},{whale2},{everyone else}. Scaling alone does NOT fix this - you must reshape the "
          "distribution. That is the log's job.")
h2(d, "What log does, and why log1p")
bullet(d, "Log compresses the tail and turns multiplicative/ratio structure into additive: each 10x jump in "
          "spend becomes the same fixed step. Distance then means 'how many times bigger', not 'how many "
          "pounds bigger' - the right notion of customer similarity.")
bullet(d, "Use log1p = log(1+x): defined at 0 (some netted Monetary or Recency=0 on the anchor day), stable "
          "for small values, ~identical to log for large.")
h2(d, "Decision & verification")
bullet(d, "Apply log1p to Frequency, Monetary, AvgBasket. Leave Recency raw unless its skewness says otherwise "
          "(it is usually mild / sometimes bimodal) - do not auto-log every feature.")
bullet(d, "Verify, do not assert: histograms before/after, Fisher skewness pre/post (target near 0), Q-Q plot. "
          "Optional rigor: run Box-Cox and report the optimal lambda lands near 0, empirically justifying log.")

h1(d, "Issue 2 - Scaling")
para(d, "Log fixes each feature's SHAPE; features still live on different ranges (Recency in hundreds of "
        "days vs log features in single digits to ~12), so distance would now be dominated by Recency. "
        "Scaling makes features contribute comparably. Order matters: LOG first (reshape), THEN scale "
        "(re-range) - scaling is affine and cannot fix skew.")
table(d, ["Scaler", "Center / Spread", "Outlier-sensitive?", "Verdict"], [
    ["StandardScaler", "mean / std", "Yes (mean & std pulled by extremes)", "Run as experiment / robustness check"],
    ["RobustScaler", "median / IQR", "No (middle 50% only)", "CHOSEN primary - data is part-wholesale"],
    ["MinMaxScaler", "min / max-min", "Very", "Rejected - one whale squashes everyone"],
    ["Normalizer (L2)", "per-row norm", "n/a", "Rejected - rescales customers, destroys magnitude"],
])
bullet(d, "Decision: log1p -> RobustScaler as primary; ALSO run StandardScaler and compare whether segments / "
          "validation metrics are stable across the two (a reportable robustness result).")
bullet(d, "Standardising to equal spread is an IMPLICIT decision that all features are equally important - "
          "name it explicitly. (See Issue 3.)")
bullet(d, "Hygiene: fit the scaler once on the full feature table and persist it (joblib); same scaled matrix "
          "feeds K-Means and the PCA viz. This scaler is for clustering only - the CLV holdout uses raw "
          "counts/dates via lifetimes, a separate pipeline.")

h1(d, "Issue 3 - Weighting & redundancy (do NOT just default to equal)")
para(d, "'Equal weight per column' only equals 'fair' if columns are independent constructs. Ours are not.")
h2(d, "Two redundancy problems")
bullet(d, "AvgBasket is a DETERMINISTIC function of the others: AvgBasket = Monetary / Frequency, so in log "
          "space log AvgBasket = log Monetary - log Frequency (exact linear combination = perfect "
          "multicollinearity). Feeding all three adds zero new information and lets the 'spend' direction "
          "count extra.")
bullet(d, "Even otherwise, F, M and Tenure are correlated - much of their variance is one underlying "
          "'how much of a customer are you' axis. Equal-per-column then TRIPLE-COUNTS that axis and drowns "
          "out Recency, which under framing A is the most decision-relevant feature (churn signal).")
h2(d, "The three philosophies")
table(d, ["Option", "Idea", "Trade-off"], [
    ["A. Fix redundancy, then equal", "Drop deterministic/near-duplicate features, then weight equally",
     "Interpretable baseline; CHOSEN starting point"],
    ["B. Business weighting", "Up-weight Recency & Monetary (the decision drivers)",
     "Defensible but subjective - 'how did you pick weights?'"],
    ["C. De-correlate (PCA-whiten / Mahalanobis)", "Let covariance handle correlated features automatically",
     "Principled, but components are harder to interpret in RFM terms"],
])
h2(d, "Decided path")
numbered(d, "Diagnose first: compute and SHOW the correlation matrix of the (log) features - let data reveal "
            "redundancy, do not just assert it. (Observability artifact.)")
numbered(d, "Remove the deterministic redundancy (AvgBasket = M/F) - a correctness fix, not a preference. "
            "Plan: make AvgBasket profiling-only; cluster on Recency / Frequency / Monetary / Tenure "
            "(re-check whether Tenure is largely redundant with Frequency and drop if so).")
numbered(d, "Default to equal weighting on the cleaned, non-redundant set as the interpretable baseline.")
numbered(d, "Run as tracked experiments: (i) business-weighted (R & M up), (ii) PCA-whitened / Mahalanobis; "
            "compare segments + validation metrics across all three. Equal weighting is EARNED by showing "
            "redundancy is handled and weighting choice does not secretly drive the result.")
save(d, "09_Feature_Treatment.docx")


# ---------------------------------------------------------------------------
# 10 - CHOOSING & VALIDATING K  (Design discussion #4, 18 June 2026)
# ---------------------------------------------------------------------------
d = new_doc("Choosing & Validating K", "Design discussion #4 - how many segments, and proving they are real")
para(d, "Added 18 June 2026. The heart of the project's rigor. Deciding the number of segments K is a "
        "JUDGEMENT triangulated from evidence across three legs - not a number read off one chart. Standing "
        "rule: every metric below is produced as a comparison artifact (observability).")

h1(d, "The three-leg framework (decide K at the intersection, not on one leg)")
table(d, ["Leg", "Question it answers", "Tools"], [
    ["Statistical", "How many blobs, geometrically?", "silhouette + gap (primary); elbow + CH/DB (corroborate)"],
    ["Stability", "Does that K reproduce, or is it one lucky run?", "bootstrap Jaccard, consensus matrix, seed ARI"],
    ["Business", "Is that K actionable?", "one nameable marketing action per segment"],
])
para(d, "Internal indices alone are the weakest single basis - they only NARROW K. The decision is made by "
        "weighing all three legs.")

h1(d, "The honest truth about this data")
para(d, "Customer behaviour is a CONTINUUM, not discrete islands - people spread smoothly from one-time "
        "bargain buyer to loyal whale. There may be no single 'true' K, and the indices will often disagree "
        "precisely because there is no clean answer. So the framing is not 'discover the correct number of "
        "segments' but 'choose a USEFUL number that is statistically supported and stable.' Stating this out "
        "loud is the mature position.")

h1(d, "(A) Internal indices for choosing K")
h2(d, "Elbow / inertia (intuition only - insufficient)")
bullet(d, "Inertia = within-cluster sum of squares. It ALWAYS falls as K rises and hits 0 at K=n, so you "
          "cannot minimise it - you look for the 'bend' where gains flatten. Subjective; two analysts read "
          "different elbows.")
bullet(d, "Make it objective with the Kneedle algorithm (point of maximum curvature). Use for intuition and "
          "to narrow the range, never as sole justification. (It is the clustering analogue of a PCA scree plot.)")
h2(d, "Silhouette (the workhorse)")
bullet(d, "Per point: s = (b - a) / max(a, b), where a = mean distance to own-cluster points, b = mean "
          "distance to the nearest OTHER cluster. s~+1 cozy inside; s~0 on a border; s<0 likely mis-assigned. "
          "Average over points -> one score per K; pick the K that maximises it.")
bullet(d, "The per-cluster silhouette PLOT shows WHICH clusters are tight vs mushy - a quality diagnostic, "
          "not just a count.")
bullet(d, "Rule of thumb (Kaufman & Rousseeuw): >0.7 strong, 0.5-0.7 reasonable, 0.25-0.5 weak, <0.25 no real "
          "structure. For behavioural RFM, 0.3-0.5 is NORMAL and honest - do not chase 0.7 (forcing it = "
          "over-splitting).")
bullet(d, "Costs O(n^2) (sample if it bites); implicitly favours convex, equal-size clusters.")
h2(d, "Gap statistic (the principled one)")
bullet(d, "Compares your clustering's tightness to clustering pure RANDOM data with the same bounding box: "
          "Gap(K) = E_null[log W_K] - log W_K. It corrects for the fact that inertia always drops.")
bullet(d, "Selection rule: smallest K with Gap(K) >= Gap(K+1) - s_(K+1) (a built-in Occam's razor; s is the "
          "reference standard error).")
bullet(d, "Two nulls: a plain uniform box, or a PCA-aligned box (preferred, more conservative). Uniquely, the "
          "gap statistic CAN return K=1 (no real clusters) - essential for an honesty-first project.")
h2(d, "Calinski-Harabasz & Davies-Bouldin (fast corroboration)")
bullet(d, "CH = variance ratio (between / within), F-statistic-like; HIGHER better. DB = average worst-case "
          "cluster-pair similarity; LOWER better. Both favour spherical clusters (CH shares K-Means' "
          "assumptions) -> use as corroboration, not the verdict.")
h2(d, "Why triangulate")
para(d, "Each method is biased differently (elbow subjective; silhouette/CH/DB favour spheres; gap depends on "
        "the null). They fail in DIFFERENT directions, so agreement among them is robust precisely because "
        "their biases do not coincide. Disagreement is itself a finding (usually gradient structure).")

h1(d, "(B) Stability - proving the segments are REAL")
para(d, "Internal indices say 'how many blobs' but not whether the solution reproduces. A clustering can "
        "score well yet be fragile. Principle: real structure reproduces under perturbation; artifacts do not.")
h2(d, "Bootstrap cluster stability (Hennig clusterboot / Jaccard)")
bullet(d, "Subsample customers (prefer 80% WITHOUT replacement), re-cluster, and for each ORIGINAL cluster "
          "find its best-matching bootstrap cluster via Jaccard overlap = |A and B| / |A or B|. Average over "
          "~100 runs.")
bullet(d, "Rule of thumb: mean Jaccard >0.85 highly stable; 0.75-0.85 stable; 0.60-0.75 suggestive; <0.60 "
          "DISSOLVES (not a real segment). Scores EACH cluster - so you can find '3 solid + 1 dissolving', "
          "which usually means K is one too high.")
h2(d, "Consensus clustering (Monti et al.)")
bullet(d, "Build an n x n matrix = fraction of runs in which each customer PAIR co-clusters. For the right K "
          "it is near block-diagonal (entries near 0 or 1). Visualise as a heatmap (crisp blocks = stable) "
          "and quantify crispness - an independent SECOND way to pick K.")
h2(d, "Seed / initialisation stability (cheap smell test)")
bullet(d, "Run many random inits; check they converge (high pairwise ARI, low inertia variance). If different "
          "seeds give wildly different clusters, structure is weak regardless of silhouette. (This is what "
          "n_init does internally.)")
h2(d, "Comparing partitions - the label-switching problem")
bullet(d, "You cannot match label NUMBERS ('cluster 2' is arbitrary). Use label-invariant measures: ARI "
          "(whole partition, chance-corrected, 0=random 1=identical) for run-to-run agreement; Jaccard for a "
          "single cluster's survival. For different sampled points, compare on common points or assign by "
          "nearest centroid.")
h2(d, "Pitfalls")
bullet(d, "Bootstrap WITH replacement creates duplicate points (distance 0) -> degenerate micro-clusters; "
          "prefer subsampling.")
bullet(d, "Stability is NECESSARY not SUFFICIENT: a trivially separated wrong solution can be very stable, and "
          "stability tends to favour smaller K. Read it WITH internal validity + business sense, never alone.")

h1(d, "Decision procedure")
numbered(d, "Search K = 2-8 (a sane business range - nobody runs 12 playbooks).")
numbered(d, "Lead with silhouette + gap statistic; corroborate with elbow + CH/DB.")
numbered(d, "Keep the small set of CANDIDATE K's the statistics support (likely 3-5); do not pick on a third "
            "decimal of one index.")
numbered(d, "Break the tie with stability (every segment must survive resampling) + interpretability.")
numbered(d, "Document the call; disagreement among indices is a FINDING, not an embarrassment.")
para(d, "Prediction: this lands at 3-5 segments - where RFM behavioural segmentation usually settles, both "
        "statistically and because that is how many distinct playbooks a marketing team can run. Avoid: "
        "cherry-picking the prettiest-K index; over-trusting one number; chasing silhouette 0.7 on "
        "behavioural data.")
save(d, "10_Choosing_and_Validating_K.docx")


# ---------------------------------------------------------------------------
# 11 - CLUSTERING METHOD COMPARISON  (Design discussion #5, 18 June 2026)
# ---------------------------------------------------------------------------
d = new_doc("Clustering Method Comparison", "Design discussion #5 - K-Means vs GMM vs hierarchical")
para(d, "Added 18 June 2026. We do NOT just run K-Means. Every algorithm imposes a shape assumption and will "
        "return clusters even if that assumption is wrong, so a single method's output is confounded (real "
        "structure vs method artifact). Running methods with DIFFERENT assumptions and checking whether they "
        "AGREE is how we separate real from artifact.")

h1(d, "The interview trap: 'your data is a continuum - why K-Means?'")
para(d, "Strongest answer: we are NOT claiming hidden discrete groups (it is a gradient). We use clustering "
        "to impose a useful, low-distortion PARTITION on that continuum for a business that can run only a "
        "handful of playbooks. K-Means is vector quantization - it finds the K prototype customers that "
        "summarise the distribution with minimum information loss - which is exactly 'pick a few "
        "representative segments.' And we did not assume it: we benchmarked against GMM and hierarchical.")
h2(d, "Separate two questions (the interviewer conflates them)")
bullet(d, "'Does the data have natural discrete clusters?' - often NO (continuum).")
bullet(d, "'Is a K-way partition a useful, low-distortion summary?' - often YES. Segmentation is the second "
          "question. Age brackets, tax bands, and credit-score bands all partition continua usefully despite "
          "no natural gaps.")
h2(d, "The caveat to VOLUNTEER (volunteering it scores the points)")
para(d, "Boundary customers are genuinely ambiguous on a continuum, so we do not over-claim crisp groups. "
        "That is exactly why we report the silhouette (flags boundary/negative points), cross-check the GMM's "
        "soft probabilities, and treat segments as a deliberate discretisation rather than a discovery of "
        "natural kinds. Naming the limitation before being asked signals maturity.")

h1(d, "K-Means - the interpretable quantizer")
bullet(d, "Minimises within-cluster squared Euclidean distance; carves space into Voronoi cells with hard "
          "straight boundaries. Assumes spherical, similar-size, similar-variance clusters; hard assignment.")
bullet(d, "Strengths: fast, scalable, INTERPRETABLE centroids (each = a describable average customer). The "
          "optimal quantizer = our continuum justification. Weaknesses: spherical assumption often wrong; "
          "hard boundaries poor for a continuum; must pre-specify K. Role: deployed baseline.")

h1(d, "GMM - honest about overlap (plain + rigorous)")
para(d, "Plain: instead of hard bubbles, GMM says the customer base is a BLEND of a few 'types' and each "
        "person is a mix (e.g. Erin = 60% loyal-regular, 40% lapsed-spender). Soft percentages, not a forced box.")
bullet(d, "Generative model: a mixture of K Gaussians, each with its own MEAN (archetype) and COVARIANCE "
          "(shape/size/tilt of the oval) plus a mixing weight (segment size). K-Means bubbles must be round "
          "and equal; GMM ovals can stretch, tilt, and differ in size.")
bullet(d, "Fit by EM: responsibility = posterior probability a component generated a customer (the soft "
          "membership). E-step computes responsibilities; M-step re-estimates each Gaussian by "
          "responsibility-weighted averages; iterate. Likelihood increases each round but only to a LOCAL "
          "optimum -> several inits, usually k-means-initialised (so K-Means is a precursor, not a rival).")
bullet(d, "covariance_type (spherical / diag / tied / full) = how fancy each oval may be; a bias-variance "
          "dial. spherical ~ soft K-Means; full = flexible but overfits. Chosen by BIC, not by eye.")
bullet(d, "Payoff: soft memberships DETECT fence-sitters (entropy of the membership vector, or 1 - max prob). "
          "You can report 'X% of customers confidently assigned vs on a boundary' - a direct, honest "
          "statement about the continuum.")
bullet(d, "Theory link: K-Means is GMM in a limit (spherical covariance, variance -> 0, hard assignment) -> "
          "GMM strictly GENERALISES K-Means.")
bullet(d, "Failure modes: singularities (a component collapses onto a point, variance -> 0, likelihood -> "
          "infinity) -> regularise (reg_covar); non-Gaussianity (RFM not perfectly Gaussian even after log) -> "
          "GMM may add SPURIOUS components just to patch shape, so do not over-read 'BIC wants 6'; local "
          "optima -> multiple inits.")

h1(d, "Hierarchical / agglomerative - the structure visualiser (plain + rigorous)")
para(d, "Plain: build a family tree by repeatedly merging the two most-similar groups until one remains, then "
        "SLICE the tree wherever you like to get segments. Anna+Beth (near-identical) merge first/low; Carl "
        "(giant rare order) is absorbed last/high.")
bullet(d, "Mechanics: start with n singletons -> merge the closest pair (by linkage) -> update distances "
          "(Lance-Williams recurrence) -> record the dendrogram. Cut at a height / at K leaves to get the "
          "partition.")
bullet(d, "Linkage = how 'distance between groups' is defined: single (closest members; chains), complete "
          "(farthest; compact), average (mean; balanced), Ward (merge that adds least within-cluster "
          "variance; K-Means-like). Default: Ward + Euclidean for comparability.")
bullet(d, "Reading the dendrogram = the continuum answer made visual: a long tall gap before a merge = a "
          "natural split (cut it); steady merges with no gaps = a GRADIENT (the cut is your choice, not a "
          "discovery). You can show the tree, mark the cut, and justify K visually.")
bullet(d, "Cophenetic correlation = corr(original pairwise distances, tree-implied distances); >0.7 means the "
          "tree faithfully represents the data - a one-number 'is this dendrogram honest?' check.")
bullet(d, "Costs: O(n^2) memory for the distance matrix (~5,900 customers -> a few hundred MB; fine on a "
          "laptop, breaks first at scale); greedy and irreversible; centroid/median linkage can 'invert' -> "
          "prefer Ward.")

h1(d, "What we are NOT using - DBSCAN")
para(d, "Density-based clustering needs dense regions separated by SPARSE GAPS. A continuum has no gaps, so "
        "DBSCAN collapses to one blob (as the original notebook found). Knowing WHY it fails is understanding, "
        "not a gap.")

h1(d, "The comparison machinery")
h2(d, "ARI - do two groupings agree?")
para(d, "Over all PAIRS of customers, count those the two methods agree on (both-together or both-apart), "
        "then adjust so random scores 0. ARI = 1 identical, 0 chance-level, <0 worse than random. "
        "Label-invariant. Headline use: a 3x3 ARI matrix among K-Means / GMM / Ward at the chosen K - high "
        "off-diagonal (>=0.7) means three differently-biased methods agree on who-belongs-with-whom -> the "
        "segmentation is REAL, not an artifact of one method.")
h2(d, "BIC - the fair count of segments (GMM only)")
para(d, "GMM is a likelihood model, so BIC = -2 ln L + p ln n (p = free params, grows with K and covariance "
        "type). Lower is better; its ln n penalty is harsher than AIC's, so BIC prefers simpler models and "
        "guards against spurious components. Plot BIC vs K per covariance type -> a K vote from a DIFFERENT "
        "family of evidence than silhouette/gap.")
h2(d, "Crowning a method fairly (avoid circularity)")
para(d, "Do not judge with a referee that shares a method's assumptions: silhouette/CH/Ward all favour "
        "spheres and so tilt toward K-Means/Ward. Weigh a BASKET: internal indices + stability + GMM BIC + "
        "cross-method ARI + interpretability.")

h1(d, "Likely decision (let the evidence confirm)")
bullet(d, "K-Means = deployed primary (interpretable centroids, scalable).")
bullet(d, "GMM = continuum-aware cross-check (soft boundary memberships + BIC corroboration of K).")
bullet(d, "Hierarchical (Ward) = visual/structural corroboration (dendrogram narrative + cophenetic check).")
bullet(d, "Headline credibility result = the cross-method ARI agreement. Follow the evidence if GMM materially "
          "wins on BIC/validation.")
para(d, "Observability deliverables: method x metric table, 3x3 ARI matrix, BIC-vs-K curves, the dendrogram "
        "with its cut, and PCA scatter of each method's labels overlaid.")
save(d, "11_Clustering_Method_Comparison.docx")


# ---------------------------------------------------------------------------
# 12 - CLV AND SEGMENTS INTEGRATION  (Design discussion #6, 18 June 2026)
# ---------------------------------------------------------------------------
d = new_doc("CLV and Segments Integration", "Design discussion #6 - how the value model and the clusters fit together")
para(d, "Added 18 June 2026. We build TWO models - a clustering (groups customers by behaviour) and a CLV "
        "model (BG/NBD + Gamma-Gamma, predicts future value). This records how they relate.")

h1(d, "The options")
table(d, ["Option", "Idea", "Verdict"], [
    ["CLV as a clustering input", "Feed predicted CLV in as another feature", "Rejected - circular (see below)"],
    ["CLV as a post-hoc layer", "Cluster on behaviour, attach CLV per customer afterward", "CHOSEN"],
    ["One combined model", "A single joint model of behaviour + value", "Rejected - overcomplicated, loses interpretability"],
])

h1(d, "Why CLV must NOT be a clustering input")
bullet(d, "Circularity: CLV is built FROM Recency/Frequency/Monetary - the same inputs the clustering already "
          "uses - so feeding it back double-counts the spend signal (the redundancy trap from discussion #3 "
          "in a new costume).")
bullet(d, "Interpretability: we want segments defined by BEHAVIOUR, then to state what each behavioural segment "
          "is worth. If CLV helps define the groups, we can no longer say 'this behaviour group turns out to "
          "be worth the most' - value already shaped the group.")
bullet(d, "Keeping them separate preserves a free validation (next section) that merging would destroy.")

h1(d, "How they meet: a join on Customer ID")
para(d, "Two separate pipelines that meet at the customer level:")
bullet(d, "Clustering pipeline: scaled log-RFM + tenure -> segment label.")
bullet(d, "CLV pipeline: RAW frequency / recency / tenure / monetary -> lifetimes model -> predicted future "
          "value. Note: CLV uses raw counts/dates - NOT scaled or logged; a different pipeline entirely.")
para(d, "Each customer ends with two attributes - a segment and a predicted CLV - joined into one final table.")

h1(d, "The free validation (internal consistency check)")
para(d, "Two INDEPENDENT methods both point at 'who is valuable': clustering ('these form the Champion "
        "behaviour group') and CLV ('these have the highest predicted future value'). If high-CLV customers "
        "concentrate in the Champion segment, the two independent methods AGREE -> strong evidence both "
        "capture something real. If they disagree (e.g. an active-looking segment with low CLV -> possibly "
        "about to churn), that is itself a finding. Deliverable: a segment x CLV summary table (observability).")

h1(d, "The business deliverable: the segment x CLV action grid")
para(d, "Combine segment (behaviour / lifecycle) with CLV (value) into a prioritised action - this is where "
        "framing A pays off: spend follows expected RETURN (CLV) WITHIN behavioural context (segment).")
table(d, ["Segment \\ Value", "High CLV", "Low CLV"], [
    ["At-risk", "TOP priority win-back (valuable AND leaving)", "Let them lapse"],
    ["Champion", "Protect / VIP perks", "Steady nurture"],
    ["New", "Invest to grow", "Light-touch onboarding"],
])
para(d, "A high-value customer who is slipping is the single best place to spend a retention pound - and only "
        "the segment x value combination surfaces them.")

h1(d, "Sub-decisions inside CLV")
bullet(d, "Horizon: predict CLV over a fixed future window - default 12 months (clean business-planning unit).")
bullet(d, "Discounting: optionally apply DCF (a pound next year is worth less than today) - a rigor touch; "
          "can be a v2.")
bullet(d, "Bonus signal: BG/NBD also outputs P(alive), a per-customer churn-risk probability that complements "
          "the segment's lifecycle read.")
save(d, "12_CLV_and_Segments_Integration.docx")


# ---------------------------------------------------------------------------
# 13 - SEGMENT PROFILING & VALIDATION  (Design discussion #7a, 18 June 2026)
# ---------------------------------------------------------------------------
d = new_doc("Segment Profiling & Validation", "Design discussion #7a - turning clusters into named, validated personas")
para(d, "Added 18 June 2026. Once clustering produces labels, this is how we turn them into named personas "
        "(Section 1) and PROVE the personas are real and distinct (Section 2). Companion doc 14 covers "
        "recommendations + success criteria.")

h1(d, "Section 1 - From clusters to personas (profiling discipline)")
para(d, "A cluster label (0,1,2,3) is meaningless until characterised. Two rules.")
h2(d, "Rule (a) - profile on ORIGINAL units, not the modelling space")
para(d, "Clustering ran on scaled log-features; those centroids are uninterpretable ('log-monetary z=0.8' "
        "means nothing to a marketer). For naming, invert back to RAW units and report per-segment MEDIANS "
        "(robust to residual skew). The segment profile table, with the business reason for each column:")
table(d, ["Column", "Business meaning (what decision it drives)"], [
    ["Segment", "Unit of action - one segment = one playbook + one budget line"],
    ["n (%)", "Size / reach - sizes the opportunity & cost; whether a dedicated playbook is worth it; flags non-viable tiny groups"],
    ["Median Recency (days)", "Churn-risk / urgency - timing of intervention; which lifecycle stage"],
    ["Median Frequency (orders)", "Loyalty / habit strength - contact cadence, cross-sell, loyalty-programme fit"],
    ["Median Monetary (GBP)", "Realised value to date - tiering; how much margin there is to fund perks"],
    ["Tenure", "Relationship maturity - context for the rest (new-and-rising vs long-loyal); sharpens churn read"],
    ["Rev share (current)", "Revenue concentration TODAY - what is at risk if the segment churns; justifies defensive spend"],
    ["CLV share (predicted)", "Future-value concentration - where value is heading; justifies growth spend"],
])
para(d, "The most decision-relevant signal is the GAP between Rev share (now) and CLV share (future):")
table(d, ["Pattern", "Reading", "Action"], [
    ["High rev share, low CLV share", "Declining segment - big today, fading", "Harvest, do not over-invest"],
    ["Low rev share, high CLV share", "Rising segment - small today, growing", "Invest ahead of the curve"],
    ["High both", "The genuine core", "Protect at all costs"],
    ["Low both", "Marginal", "Deprioritise"],
])
h2(d, "How n% is calculated")
bullet(d, "n_k = distinct CUSTOMERS in segment k, computed on the one-row-per-customer table (NOT invoice "
          "lines). %_k = n_k / N_total x 100.")
bullet(d, "Denominator N_total = the CLEANED, clustered customer base (~5,900) - not all transactions or all "
          "humans who ever transacted. State the scope explicitly: '% of analysable customers'.")
bullet(d, "Hard vs soft: K-Means/hierarchical give an unambiguous count. GMM is soft -> use the hard (argmax "
          "responsibility) count for comparability across methods; the mixing weight pi_k is the model's own "
          "prevalence view. Report hard counts in the profile table.")
h2(d, "Rule (b) - name AFTER inspecting the profile, not before")
para(d, "Derive the name from the behavioural signature; do NOT impose textbook personas "
        "(Champions/Loyal/Potential/At-risk/Hibernating) and bend the data to fit. 'Recent + frequent + "
        "high-spend' earns 'Champion'; a 'frequent but low-value wholesaler' group is named for what it "
        "actually is. Forcing canonical names onto non-matching clusters is confirmation bias - a real "
        "failure mode.")

h1(d, "Section 2 - Validating the personas are real and distinct")
para(d, "Separate from stability (#4: 'do the clusters reproduce?'). This asks: 'are the named segments "
        "genuinely different, on dimensions that matter, and would an independent method agree?'")
h2(d, "The circularity problem and the three tiers of evidence")
para(d, "Testing whether segments differ on the features used to cluster is near-circular (clustering "
        "optimises exactly that). So evidence comes in tiers:")
table(d, ["Tier", "Test", "Strength"], [
    ["1 (weak/circular)", "Separation on the CLUSTERING features", "Confirms + identifies which features; mildly circular"],
    ["2 (strong)", "Separation on EXTERNAL variables the clustering never saw", "Non-circular proof"],
    ["3 (strong)", "Agreement with an INDEPENDENT labelling (rule-based RFM)", "Corroboration"],
])
h2(d, "(A) Separation on clustering features - quantify, but read carefully")
bullet(d, "Omnibus = Kruskal-Wallis (non-parametric ANOVA on ranks) - chosen because RFM is skewed and "
          "heteroscedastic even after log, violating ANOVA's assumptions. Being rank-based, it gives IDENTICAL "
          "results on raw/log/scaled features (no 'which space?' choice). A rejection means the DISTRIBUTIONS "
          "differ (not strictly the medians, since segment shapes differ).")
bullet(d, "Large-n caveat: at ~5,900 customers everything is 'significant' (p~0); p-values are uninformative. "
          "Report EFFECT SIZE: eta^2_H = (H - k + 1)/(N - k) in [0,1] (or epsilon^2). Rank features by it -> "
          "the discriminative-feature ranking that justifies the names (usually Recency & Monetary; Tenure weak).")
bullet(d, "Post-hoc = Dunn's test (uses the pooled KW ranks) with Holm correction -> which segment PAIRS "
          "differ; pairwise magnitude via Cliff's delta. A pair small on ALL features = secret twins = K too "
          "high (feeds back to discussion #4).")
bullet(d, "Honest framing: this tier is DIAGNOSTIC (which features, which pairs collapse), not proof of 'real'.")
h2(d, "(B) Multivariate / classifier-based separation")
bullet(d, "(A) tests one feature at a time; segments can overlap on each feature ALONE yet separate in "
          "COMBINATION (e.g. high-frequency/low-basket vs low-frequency/high-basket sit in different corners).")
bullet(d, "Centroid separation: distance between segment centres in the full scaled space.")
bullet(d, "Classifier check: hide the labels, train a simple model (logistic / random forest) to predict the "
          "segment from features. High CROSS-VALIDATED accuracy = the segments are separable; the feature "
          "importances double as interpretation (which features define each segment, multivariate). Still "
          "tier-1 (same features) - CV ensures we measure real separability, not memorisation.")
h2(d, "(C) External-variable check - the GOLD STANDARD (non-circular)")
para(d, "Do the segments ALSO differ on variables the clustering never saw? If groups built only from "
        "R/F/M/Tenure also differ in what they buy or their future value, that difference is genuinely "
        "discovered, not baked in. (Analogy: sort people by height, then find the groups also differ in the "
        "sport they play - a real finding, because you did not sort on sport.)")
table(d, ["External variable", "Question", "Source (excluded from clustering)"], [
    ["Predicted CLV", "Differ in future value?", "the separate BG/NBD model"],
    ["Product mix", "Buy different things?", "StockCode / Description"],
    ["Country", "Geographic skew?", "Country"],
    ["Return rate", "Return more/less?", "cancellations / returns"],
])
bullet(d, "Numeric externals (CLV, return rate) -> effect-size comparison (now NON-circular). Categorical "
          "(product, country) -> cross-tab + chi-square / Cramer's V.")
bullet(d, "Example: segments built only from R/F/M/Tenure show avg predicted CLV of Champions 800 > Loyal 300 "
          "> At-risk 150 > One-timers 20 (GBP) - orderly differences on a variable never seen -> real (this is "
          "also the doc-12 internal-consistency check).")
bullet(d, "NULL result meaning: if segments differ on NO external variable, they may be arbitrary slices of "
          "the spend axis with no deeper meaning - report that honestly.")
bullet(d, "Interview line: 'I checked separation on predicted value, product mix, and geography - things the "
          "model never saw - not just the features I clustered on. That is evidence the segments are real "
          "customer types, not an artifact of where I drew the lines.'")
h2(d, "(Tier 3) Agreement with rule-based RFM")
bullet(d, "Build the classic RFM quintile scores (R/F/M each 1-5) -> rule-based segments, INDEPENDENT of the "
          "clustering. Cross-tab cluster vs rule segment; association via Cramer's V (or ARI). High agreement "
          "corroborates; structured divergence = the clustering REFINING the rules' arbitrary quintile cutoffs "
          "(a finding, not a fault).")
h2(d, "Visuals (observability)")
bullet(d, "Snake plot (standardised segment medians as lines - the canonical RFM visual); radar / "
          "parallel-coordinates; PCA scatter coloured by segment (note: 2-D PCA UNDERSTATES separation, so "
          "overlap there is not proof of non-separation); violin/box per discriminative feature.")
h2(d, "Pitfalls")
bullet(d, "Treating p-values as evidence at this n; claiming 'medians differ' when shapes differ; circular "
          "testing on clustering features only; over-reading 2-D PCA overlap; validating tiny segments.")
save(d, "13_Segment_Profiling_and_Validation.docx")


# ---------------------------------------------------------------------------
# 14 - RECOMMENDATIONS & SUCCESS CRITERIA  (Design discussion #7b, 18 June 2026)
# ---------------------------------------------------------------------------
d = new_doc("Recommendations & Success Criteria", "Design discussion #7b - acting on segments, and defining 'good'")
para(d, "Added 18 June 2026. Companion to doc 13. Section 3 turns validated, named segments into marketing "
        "actions; Section 4 defines what success means for an unsupervised, portfolio project.")

h1(d, "Section 3 - Recommendations / the action grid")
para(d, "We now have named, validated segments + each customer's predicted CLV. The question: what do we DO "
        "about each segment, and how much do we spend?")
h2(d, "Rules")
bullet(d, "One segment, one clear action (Protect / Grow / Nurture / Win-back / Deprioritise) - clarity is the "
          "deliverable; do not hand marketing ten vague ideas per group.")
bullet(d, "Spend follows VALUE AT STAKE (revenue share + CLV share), not headcount.")
bullet(d, "Refine with the segment x CLV grid (doc 12): within a segment, high-CLV vs low-CLV members get "
          "different intensity of the same action.")
h2(d, "Action table (example)")
table(d, ["Segment", "Action", "Why", "Spend"], [
    ["Champions (recent, high value)", "Protect - VIP perks, early access", "they'll buy anyway; do NOT discount (burns margin)", "moderate, defensive"],
    ["At-risk + high CLV", "Win-back NOW", "most recoverable value - the top-priority cell", "high per customer"],
    ["New / Potential", "Nurture to 2nd purchase", "onboarding, gentle incentive", "scaled to upside"],
    ["One-timer + low CLV", "Deprioritise", "spend > worth", "near zero"],
])
h2(d, "Value at stake beats headcount (example)")
table(d, ["Segment", "% of customers", "% of future CLV", "Priority"], [
    ["At-risk", "15%", "25%", "High - small group, big value leaking"],
    ["One-timers", "40%", "5%", "Low - huge group, little value"],
])
para(d, "Prioritising by headcount would pour budget into the 40% one-timers who are barely worth anything. "
        "Money on the line is the right lens.")
h2(d, "Optional ROI framing (only if kept honest)")
para(d, "A rough estimate ('win-back at 10% conversion recovers ~GBP X of predicted CLV for ~GBP Y spend') "
        "shows business reasoning, but MUST be flagged hypothetical - we have no campaign-response data. State "
        "the assumption out loud or it becomes the dishonesty we keep avoiding.")
h2(d, "Mistakes the grid prevents")
bullet(d, "Discounting Champions (margin lost on people who would have bought anyway).")
bullet(d, "Spending equally across segments (the whole point is UNequal spend).")
bullet(d, "Chasing low-value churned customers (sunk-cost spending).")

h1(d, "Section 4 - What 'good' looks like (success criteria)")
para(d, "Unsupervised -> there is no accuracy score. Define success deliberately, or you reach for a fake one "
        "(a high silhouette, pretty clusters). 'Good' = five things, all of which we have built toward:")
table(d, ["#", "Criterion", "Met when"], [
    ["1", "Validated", "stable (#4) + separated on external variables (2C)"],
    ["2", "Interpretable", "each segment has a clear, nameable profile in raw units"],
    ["3", "Internally consistent", "CLV concentrates in the Champion segment (doc 12)"],
    ["4", "Actionable", "every segment -> a differentiated action + quantified value at stake"],
    ["5", "Honest", "surprises & limitations reported, not hidden"],
])
h2(d, "What 'good' beats")
bullet(d, "GOOD (even if modest): 'The data is largely a continuum; only 3 robust segments survive resampling "
          "and differ on predicted CLV and product mix - here is each profile, action, and value at stake.'")
bullet(d, "BAD (even if flashy): '6 crisp, colourful segments named from a textbook' that dissolve under "
          "bootstrap, with a made-up ROI.")
para(d, "The modest-but-honest result is the SUCCESSFUL one. For a PORTFOLIO piece especially, the grade comes "
        "from the JUDGEMENT on display (validate, compare methods, admit the continuum, report what failed) - "
        "not from the cluster count or a pretty chart.")
h2(d, "Final deliverables ('done', physically)")
bullet(d, "The customer-level table: each customer -> segment + predicted CLV + recommended action.")
bullet(d, "The segment profile table (doc 13, Section 1).")
bullet(d, "The action / value-at-stake grid (Section 3).")
bullet(d, "The validation evidence (stability + separation + cross-method agreement).")
bullet(d, "A written narrative tying it all back to the problem statement.")
save(d, "14_Recommendations_and_Success_Criteria.docx")


print("\nAll documents generated in:", OUT_DIR)
